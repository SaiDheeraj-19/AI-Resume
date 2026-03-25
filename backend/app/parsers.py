import io
import re

import fitz  # PyMuPDF
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract clean text from PDF bytes using PyMuPDF."""
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text")
            text_parts.append(text)

    raw_text = "\n".join(text_parts)
    return _post_process_text(raw_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract clean text from DOCX bytes using python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    raw_text = "\n".join(text_parts)
    return _post_process_text(raw_text)


def _post_process_text(text: str) -> str:
    """
    Post-process extracted text to remove noise:
    - Remove page numbers
    - Normalize whitespace
    - Remove repeated separators
    """
    # Remove lines that look like page numbers ("Page 1 of 3", "1", "- 2 -")
    text = re.sub(r"(?m)^\s*(Page\s+\d+\s*(of\s+\d+)?|\-?\s*\d+\s*\-?)\s*$", "", text)
    # Remove long lines of dashes or equals (section separators)
    text = re.sub(r"[-=_]{3,}", " ", text)
    # Remove URLs (not needed for NLP)
    text = re.sub(r"http\S+|www\.\S+", " ", text)
    # Normalize whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
